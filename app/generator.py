from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_colored_rectangle(doc):
    """
    Add a blue colored rectangle at the top of the document spanning full width.
    """
    # Add a paragraph for the colored bar
    bar = doc.add_paragraph()
    bar_format = bar.paragraph_format
    bar_format.space_before = Pt(0)
    bar_format.space_after = Pt(0)
    bar_format.line_spacing = Pt(15)  # Half the previous height
    bar_format.left_indent = Inches(-0.5)  # Extend to left edge
    bar_format.right_indent = Inches(-0.5)  # Extend to right edge
    
    # Set shading (background color) for the paragraph
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), '156082')  # Blue color
    bar._element.get_or_add_pPr().append(shading_elm)
    
    return bar

def generate_lab_sheet(student_name, student_id, module_name, module_code, practical_number, logo_path):
    """
    Generate a lab sheet document based on the template.
    
    Args:
        student_name: Student's full name
        student_id: Student ID number
        module_name: Name of the module
        module_code: Module code
        practical_number: Practical number (e.g., "Practical 06")
        logo_path: Path to the university logo image
    """
    # Create a new Document
    doc = Document()
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
    
    # Add the blue colored rectangle at the top
    add_colored_rectangle(doc)
    
    # Add the university logo (centered from page edges, not margins)
    logo_paragraph = doc.add_paragraph()
    logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo_paragraph.paragraph_format.left_indent = Inches(-0.5)
    logo_paragraph.paragraph_format.right_indent = Inches(-0.5)
    logo_run = logo_paragraph.add_run()
    logo_run.add_picture(logo_path, width=Inches(1.1), height=Inches(1.05))
    logo_paragraph.paragraph_format.space_after = Pt(6)
    
    # Add module name and code (centered from page edges, size 20)
    module_paragraph = doc.add_paragraph()
    module_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    module_paragraph.paragraph_format.left_indent = Inches(-0.5)
    module_paragraph.paragraph_format.right_indent = Inches(-0.5)
    module_run = module_paragraph.add_run(f'{module_name} -- {module_code}')
    module_run.bold = True
    module_run.font.size = Pt(20)
    module_run.font.name = 'Times New Roman'
    module_paragraph.paragraph_format.space_after = Pt(12)  # One line space
    
    # Add practical number (size 12, left aligned)
    practical_paragraph = doc.add_paragraph()
    practical_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    practical_run = practical_paragraph.add_run(f'{practical_number}')
    practical_run.bold = True
    practical_run.font.size = Pt(12)
    practical_run.font.name = 'Aptos (Body)'
    practical_paragraph.paragraph_format.space_after = Pt(2)
    
    # Add student name and ID on next line (size 12, left aligned)
    name_paragraph = doc.add_paragraph()
    name_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    name_run = name_paragraph.add_run(f'{student_name} - {student_id}')
    name_run.bold = True
    name_run.font.size = Pt(12)
    name_run.font.name = 'Aptos (Body)'
    name_paragraph.paragraph_format.space_after = Pt(0)
    
    # Add the horizontal line (appropriate length)
    line_paragraph = doc.add_paragraph()
    line_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    line_run = line_paragraph.add_run('_' * 95)
    line_run.font.name = 'Times New Roman'
    line_paragraph.paragraph_format.space_after = Pt(12)
    
    # Save the document
    output_filename = f'{practical_number.replace(" ", "_")}_{student_id}.docx'
    doc.save(output_filename)
    print(f'Lab sheet generated successfully: {output_filename}')
    return output_filename

# Example usage (hardcoded for now)
if __name__ == '__main__':
    # Hardcoded details - will be replaced with user input later
    student_name = 'NONIS P.K.D.T.'
    student_id = 'IT23614130'
    module_name = 'Programming Paradigms'
    module_code = 'SE2052'
    practical_number = 'Practical 06'
    
    # Logo path - IMPORTANT: Place your university logo here
    # For the project structure, put the logo in: app/ui/assets/logo.png
    logo_path = 'app/ui/assets/logo.png'
    
    try:
        generate_lab_sheet(
            student_name=student_name,
            student_id=student_id,
            module_name=module_name,
            module_code=module_code,
            practical_number=practical_number,
            logo_path=logo_path
        )
    except FileNotFoundError:
        print(f'Error: Logo file not found at {logo_path}')
        print('Please place your university logo at: app/ui/assets/logo.png')
    except Exception as e:
        print(f'Error generating lab sheet: {e}')